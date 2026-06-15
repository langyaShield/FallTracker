import os
import fitz  # pymupdf
import pytesseract
from pdf2image import convert_from_path
from PIL import Image

# N-BUG-6: .docx 解析支持（python-docx 可选依赖）
try:
    from docx import Document as DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def extract_text_from_file(file_path: str, progress_callback=None) -> str:
    """Extract text from a PDF, image, or Word file.

    Strategy for PDF:
      1. Try PyMuPDF direct text extraction (instant, works for text-based PDFs)
      2. If no text found, fall back to pytesseract OCR on rendered pages

    For images: run pytesseract OCR directly.
    For .docx: use python-docx to read paragraphs (N-BUG-6).

    Args:
        file_path: Path to the file.
        progress_callback: Optional callable(current_page, total_pages, stage) for progress.
            stage is one of: "extracting", "ocr_page", "done"

    Returns:
        Extracted text string, or empty string on failure.
    """
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == '.pdf':
            return _extract_from_pdf(file_path, progress_callback)
        elif ext in ('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif'):
            return _extract_from_image(file_path, progress_callback)
        elif ext == '.docx':
            return _extract_from_docx(file_path, progress_callback)
        else:
            return ""
    except Exception as e:
        return f"[OCR失败: {str(e)}]"


def _extract_from_pdf(file_path: str, progress_callback=None) -> str:
    """Extract text from PDF: try direct extraction first, then OCR fallback."""
    doc = fitz.open(file_path)
    total_pages = len(doc)

    # Stage 1: Try direct text extraction (instant for text-based PDFs)
    if progress_callback:
        progress_callback(0, total_pages, "extracting")

    all_text = []
    for i, page in enumerate(doc):
        text = page.get_text("text").strip()
        all_text.append(text)
        if progress_callback:
            progress_callback(i + 1, total_pages, "extracting")
    doc.close()

    # If we got meaningful text, return it
    combined = "\n\n".join(t for t in all_text if t)
    if len(combined.strip()) > 20:  # threshold: enough text found
        if progress_callback:
            progress_callback(total_pages, total_pages, "done")
        return combined

    # Stage 2: No text found - likely a scanned PDF, use OCR
    if progress_callback:
        progress_callback(0, total_pages, "ocr_page")

    images = convert_from_path(file_path, dpi=200)
    ocr_texts = []
    for i, img in enumerate(images):
        page_text = pytesseract.image_to_string(img, lang='chi_sim+eng')
        ocr_texts.append(page_text.strip())
        if progress_callback:
            progress_callback(i + 1, total_pages, "ocr_page")

    if progress_callback:
        progress_callback(total_pages, total_pages, "done")

    return "\n\n".join(t for t in ocr_texts if t)


def _extract_from_image(file_path: str, progress_callback=None) -> str:
    """Run OCR directly on an image file."""
    if progress_callback:
        progress_callback(0, 1, "ocr_page")

    img = Image.open(file_path)
    text = pytesseract.image_to_string(img, lang='chi_sim+eng')

    if progress_callback:
        progress_callback(1, 1, "done")

    return text.strip()


def _extract_from_docx(file_path: str, progress_callback=None) -> str:
    """N-BUG-6: 从 .docx 文件中提取纯文本。

    提取所有段落（含表格内的段落），用空行分隔。
    若 python-docx 依赖未安装，返回明确错误信息提示用户安装。
    """
    if not _DOCX_AVAILABLE:
        return "[OCR失败: 缺少 python-docx 依赖，请执行 pip install python-docx]"

    if progress_callback:
        progress_callback(0, 1, "extracting")

    doc = DocxDocument(file_path)
    parts: list[str] = []

    # 1. 段落正文
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 2. 表格内容（按行扫描每格的段落）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)

    if progress_callback:
        progress_callback(1, 1, "done")

    return "\n\n".join(parts)
